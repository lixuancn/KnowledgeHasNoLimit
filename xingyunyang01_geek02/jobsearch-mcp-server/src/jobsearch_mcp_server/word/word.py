import os
from docx import Document


def read_word_file(file_path):
    """读取Word文档内容"""

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    file_path = os.path.abspath(file_path)
    file_ext = os.path.splitext(file_path)[1].lower()

    if file_ext == '.docx':
        return _read_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_ext}")


def _read_docx(file_path):
    """读取.docx文件"""
    try:
        doc = Document(file_path)
        full_text = []

        # 获取所有段落
        for para in doc.paragraphs:
            if para.text:
                full_text.append(para.text)

        # 获取所有表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text:
                        row_text.append(cell.text)
                if row_text:
                    full_text.append(" | ".join(row_text))

        return "\n".join(full_text)
    except Exception as e:
        raise


if __name__ == "__main__":
    file_path = "../../excluded_folders/xingyunyang01_geek02/zhangsan_resume.docx"
    content = read_word_file(file_path)
    print(content)